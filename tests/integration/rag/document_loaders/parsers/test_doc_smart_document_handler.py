# © Artur Czarnecki. All rights reserved.
# Intergrax framework – proprietary and confidential.
# Use, modification, or distribution without written permission is prohibited.

from __future__ import annotations

import pytest
from pathlib import Path

import docx
from langchain_core.documents import Document

from intergrax.rag.document_loaders.config.document_loader_config import GLOBAL_DOCUMENT_LOADER_CONFIG
from intergrax.rag.document_loaders.handlers.doc_smart_document_handler import (
    DocSmartDocumentHandler,
)

pytestmark = pytest.mark.integration


def _create_docx(path: Path) -> None:

    d = docx.Document()
    d.add_heading("Intergrax DOCX Test", level=1)
    d.add_paragraph("Hello from DOCX handler test.")
    d.save(path)


def test_doc_handler_supports_extensions():

    handler = DocSmartDocumentHandler()

    assert handler.supports("file.docx") is True
    assert handler.supports("file.doc") is True
    assert handler.supports("file.txt") is False


def test_doc_handler_confidence():

    handler = DocSmartDocumentHandler()

    assert handler.confidence("file.docx") == GLOBAL_DOCUMENT_LOADER_CONFIG.default_builtin_handler_confidence


def test_doc_handler_builds_parser():

    handler = DocSmartDocumentHandler()

    parsers = handler.build_parsers()

    assert len(parsers) == 1
    assert parsers[0].parser_id() == "doc_smart"


def test_doc_handler_loads_docx(tmp_path: Path):

    doc_path = tmp_path / "sample.docx"

    _create_docx(doc_path)

    handler = DocSmartDocumentHandler()

    docs = handler.load(str(doc_path))

    assert docs
    assert all(isinstance(d, Document) for d in docs)

    content = " ".join(d.page_content for d in docs)

    assert "Hello from DOCX handler test." in content