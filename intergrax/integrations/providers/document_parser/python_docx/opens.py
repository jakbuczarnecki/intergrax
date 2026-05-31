# © Artur Czarnecki. All rights reserved.
# Intergrax framework – proprietary and confidential.

from __future__ import annotations

from typing import Literal, Sequence

import docx
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.documents import Document

from intergrax.integrations.contracts.document_parser import ParsedDocumentFragment
from intergrax.integrations.providers.document_parser.python_docx.config import PythonDocxIntegrationConfig

EXTRACTION_STRATEGY = Literal["auto", "fulltext", "paragraphs", "headings"]


def parse_python_docx_file(config: PythonDocxIntegrationConfig, source: str) -> list[ParsedDocumentFragment]:
    strategy: EXTRACTION_STRATEGY = config.strategy
    if strategy == "auto":
        strategy = "fulltext"

    if strategy == "fulltext":
        docs = Docx2txtLoader(source).load()
    else:
        docs = _DocxParagraphLoader(source, mode=strategy).load()

    return [
        ParsedDocumentFragment(
            text=doc.page_content or "",
            metadata={"parser_backend": "python_docx", **(doc.metadata or {})},
        )
        for doc in docs
    ]


class _DocxParagraphLoader:
    def __init__(self, path: str, mode: EXTRACTION_STRATEGY = "paragraphs"):
        self.path = path
        self.mode = mode

    def _is_heading(self, para) -> tuple[bool, int]:
        style = getattr(para.style, "name", "") or ""
        if not style:
            return (False, 0)
        lowered = style.lower()
        if lowered.startswith("heading"):
            for level in range(1, 10):
                if lowered == f"heading {level}":
                    return (True, level)
            return (True, 1)
        return (False, 0)

    def load(self) -> list[Document]:
        document = docx.Document(self.path)
        items: list[Document] = []
        heading_stack: list[str] = []
        section_ix = 0
        para_ix = 0
        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            is_head, level = self._is_heading(paragraph)
            if is_head:
                heading_stack = heading_stack[: max(level - 1, 0)]
                heading_stack.append(text)
                section_ix += 1
                if self.mode == "headings":
                    items.append(
                        Document(
                            page_content=text,
                            metadata={
                                "doc_type": "docx",
                                "source_path": self.path,
                                "section_ix": section_ix,
                                "heading_path": " / ".join(heading_stack),
                                "is_heading": True,
                            },
                        )
                    )
                continue
            if self.mode == "paragraphs":
                para_ix += 1
                items.append(
                    Document(
                        page_content=text,
                        metadata={
                            "doc_type": "docx",
                            "source_path": self.path,
                            "section_ix": section_ix,
                            "para_ix": para_ix,
                            "heading_path": " / ".join(heading_stack),
                            "is_heading": False,
                        },
                    )
                )
        return items
