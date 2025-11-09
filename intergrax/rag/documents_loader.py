# © Artur Czarnecki. All rights reserved.
# Integrax framework – proprietary and confidential.
# Use, modification, or distribution without written permission is prohibited.

import hashlib
import json
import docx
import logging
import os
import io
import base64
import requests

from intergrax.multimedia.video_loader import (
    transcribe_to_vtt,
    extract_frames_and_metadata,
)
from intergrax.multimedia.audio_loader import (
    yt_download_audio, 
    translate_audio
)

# LLM adaptery frameworka (model + provider w adapterze)
from intergrax.llm.llm_adapters import (
    LLMAdapter,
    LangChainOllamaAdapter,
)

from pathlib import Path

from pandas import DataFrame

from typing import (
    Callable, Dict, Iterable, List, Mapping, Optional, Sequence, Union,
)

from langchain_community.document_loaders import (
    Docx2txtLoader,
    PyMuPDFLoader,
    TextLoader,
    UnstructuredHTMLLoader,
)
from langchain_core.documents import Document
from tqdm.auto import tqdm
from typing import Literal

try:
    import fitz  # PyMuPDF (used only for raster OCR)
except Exception:
    fitz = None

try:
    import pytesseract
except Exception:
    pytesseract = None

from PIL import Image, ExifTags


try:
    import pandas as pd
except Exception:
    pd = None

try:
    import openpyxl  # engine for .xlsx (used by pandas)
except Exception:
    openpyxl = None

try:
    import xlrd  # engine for .xls (legacy Excel files)
except Exception:
    xlrd = None

try:
    import pytesseract
except Exception:
    pytesseract = None


logger = logging.getLogger(__name__)

MetadataFn = Callable[[Document, Path], Optional[Dict]]

DOCX_MODE = Literal["fulltext", "paragraphs", "headings"]

EXCEL_MODE = Literal["rows", "sheets", "markdown"]

class IntergraxDocumentsLoader:
    """Robust, extensible document loader with metadata injection and safety guards."""

    def __init__(
        self,
        *,
        verbose: bool = False,
        file_patterns: Iterable[str] = ("**/*",),  # include files without extension too
        extensions_map: Optional[Mapping[str, Callable[[str], object]]] = None,
        exclude_globs: Optional[Iterable[str]] = None,
        follow_symlinks: bool = False,
        max_files: Optional[int] = None,
        max_file_size_mb: Optional[int] = 64,
        docx_mode: DOCX_MODE = "fulltext",
        pdf_enable_ocr: bool = False,
        pdf_ocr_lang: str = "eng",
        pdf_ocr_dpi: int = 200,
        pdf_ocr_psm: Optional[int] = None,
        pdf_ocr_oem: Optional[int] = None,
        pdf_ocr_max_pages: Optional[int] = None,
        excel_mode: EXCEL_MODE = "rows",           # "rows" | "sheets" | "markdown"
        excel_header: int = 0,
        excel_sheet: str | int | None = None,
        excel_na_filter: bool = True,
        excel_max_rows_per_sheet: Optional[int] = None,
        csv_encoding: Optional[str] = None,
        csv_delimiter: Optional[str] = None,
        # --- image (existing) ---
        image_ocr_lang: str = "eng",
        image_ocr_psm: Optional[int] = None,
        image_ocr_oem: Optional[int] = None,
        image_extract_exif: bool = True,
        image_max_dim: Optional[int] = 2000,
        # --- image NEW: captioning modes via framework adapter ---
        image_text_mode: Literal["ocr", "caption", "both"] = "both",
        image_caption_llm: Optional[LLMAdapter] = None, 
        image_both_joiner: str = "\n\n---\n\n",
    ):
        self._verbose = verbose
        self._file_patterns = tuple(file_patterns)
        self._exclude_globs = tuple(exclude_globs or ())
        self._follow_symlinks = follow_symlinks
        self._max_files = max_files
        self._max_file_size_mb = max_file_size_mb
        self._docx_mode = docx_mode 

        self._pdf_enable_ocr = bool(pdf_enable_ocr)
        self._pdf_ocr_lang = pdf_ocr_lang
        self._pdf_ocr_dpi = int(pdf_ocr_dpi)
        self._pdf_ocr_psm = pdf_ocr_psm
        self._pdf_ocr_oem = pdf_ocr_oem
        self._pdf_ocr_max_pages = pdf_ocr_max_pages

        self._excel_mode = excel_mode
        self._excel_header = excel_header
        self._excel_sheet = excel_sheet
        self._excel_na_filter = excel_na_filter
        self._excel_max_rows_per_sheet = excel_max_rows_per_sheet
        self._csv_encoding = csv_encoding
        self._csv_delimiter = csv_delimiter

        # images (existing)
        self._image_ocr_lang = image_ocr_lang
        self._image_ocr_psm = image_ocr_psm
        self._image_ocr_oem = image_ocr_oem
        self._image_extract_exif = bool(image_extract_exif)
        self._image_max_dim = image_max_dim
        # images NEW (captioning via adapter)
        self._image_text_mode = image_text_mode
        self._image_caption_llm = image_caption_llm
        self._image_both_joiner = image_both_joiner


        default_map: Dict[str, Callable[[str], object]] = {
            ".txt":  lambda p: TextLoader(p, autodetect_encoding=True),  # autodetect
            ".docx": lambda p: Docx2txtLoader(p),            
            ".htm":  lambda p: UnstructuredHTMLLoader(p),
            ".html": lambda p: UnstructuredHTMLLoader(p),
            ".pdf":  lambda p: PdfSmartLoader(
                p,
                enable_ocr=self._pdf_enable_ocr,
                ocr_lang=self._pdf_ocr_lang,
                ocr_dpi=self._pdf_ocr_dpi,
                ocr_psm=self._pdf_ocr_psm,
                ocr_oem=self._pdf_ocr_oem,
                ocr_max_pages=self._pdf_ocr_max_pages,
            ),
            ".xlsx": lambda p: ExcelSmartLoader(
                p,
                mode=self._excel_mode,
                header=self._excel_header,
                sheet=self._excel_sheet,
                na_filter=self._excel_na_filter,
                max_rows_per_sheet=self._excel_max_rows_per_sheet,
            ),
            ".xls":  lambda p: ExcelSmartLoader(
                p,
                mode=self._excel_mode,
                header=self._excel_header,
                sheet=self._excel_sheet,
                na_filter=self._excel_na_filter,
                max_rows_per_sheet=self._excel_max_rows_per_sheet,
            ),
            ".csv":  lambda p: ExcelSmartLoader(
                p,
                mode=self._excel_mode,
                header=self._excel_header,
                na_filter=self._excel_na_filter,
                max_rows_per_sheet=self._excel_max_rows_per_sheet,
                encoding=self._csv_encoding,
                delimiter=self._csv_delimiter or ",",
            ),
            ".tsv":  lambda p: ExcelSmartLoader(
                p,
                mode=self._excel_mode,
                header=self._excel_header,
                na_filter=self._excel_na_filter,
                max_rows_per_sheet=self._excel_max_rows_per_sheet,
                encoding=self._csv_encoding,
                delimiter=self._csv_delimiter or "\t",
            ),
        }

        # images
        image_exts = (
            ".jpg",
            ".jpeg",
            ".png",
            ".tiff",
            ".bmp",
            ".webp",
            ".heic",  # optional modern iPhone format
            ".heif",  # optional, same family
        )

        for ext in image_exts:
            default_map[ext] = lambda p, _ext=ext: ImageSmartLoader(
                p,
                ocr_lang=self._image_ocr_lang,
                ocr_psm=self._image_ocr_psm,
                ocr_oem=self._image_ocr_oem,
                extract_exif=self._image_extract_exif,
                max_image_dim=self._image_max_dim,                
                text_mode=self._image_text_mode,                # "ocr" | "caption" | "both"
                caption_llm=self._image_caption_llm,
                both_joiner=self._image_both_joiner,
            )

        # videos
        video_exts = (
            ".mp4",
            ".mkv",
            ".mov",
            ".avi",
            ".webm",
            ".m4v",
            ".flv",
            ".wmv",
            ".ts",
            ".3gp",
            ".ogv",
        )

        for ext in video_exts:
            default_map[ext] = lambda p, _ext=ext: VideoSmartLoader(
                p,
                out_dir=None,                # optional, saves frames/metadata next to the video
                frames_subdir="frames",
                meta_subdir="video_meta",
                transcribe_if_missing=True,  # automatically generates VTT if missing
                whisper_model_size="base",
                whisper_language=None,       # e.g. "pl" if you want to force Polish
                frame_target_height=350,
            )


        # Audio
        audio_exts = (
            ".wav",
            ".mp3",
            ".m4a",
            ".flac",
            ".ogg",
            ".opus",
            ".aac",
            ".wma",
            ".aiff",  # aka .aif
            ".aif",
            ".mka",   # Matroska audio
        )

        for ext in audio_exts:
            default_map[ext] = lambda p, _ext=ext: AudioSmartLoader(
                path=p,
                out_dir=None,
                audio_format=_ext.lstrip("."),
                whisper_model="medium",
                whisper_language=None,
                translate=True,
            )


        if self._docx_mode == "fulltext":
            # original behavior: single Document with full text
            default_map[".docx"] = lambda p: Docx2txtLoader(p)
        elif self._docx_mode in ("paragraphs", "headings"):
            # custom loader returning a list of Documents per paragraph / heading
            default_map[".docx"] = lambda p: DocxParagraphLoader(p, mode=self._docx_mode)
        else:
            raise ValueError("docx_mode must be one of: 'fulltext', 'paragraphs', 'headings'")


        self._extensions_map: Dict[str, Callable[[str], object]] = dict(default_map)

        if extensions_map:
            for k, v in extensions_map.items():
                if not callable(v):
                    raise TypeError(f"extensions_map['{k}'] must be callable(path)->Loader")
                self._extensions_map[k.lower()] = v

        self._allowed_exts = set(self._extensions_map.keys())

    def _is_within_limits(self, file_path: Path) -> bool:
        # size guard
        if self._max_file_size_mb is not None:
            try:
                size_mb = file_path.stat().st_size / (1024 * 1024)
                if size_mb > self._max_file_size_mb:
                    if self._verbose:
                        logger.warning("[intergraxDocumentsLoader] Skipping large file (%.1f MB): %s", size_mb, file_path)
                    return False
            except OSError as e:
                logger.warning("[intergraxDocumentsLoader] Could not stat file %s: %s", file_path, e)
                return False
        return True

    def _is_excluded(self, file_path: Path, root: Path) -> bool:
        # apply exclude globs relative to root
        if not self._exclude_globs:
            return False
        rel = file_path.relative_to(root)
        for pat in self._exclude_globs:
            if rel.match(pat):
                return True
        return False

    @staticmethod
    def _stable_parent_id(path: Path) -> str:
        # stable id from absolute path
        return hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:16]


    def load_document(
        self,
        file_path: str,
        *,
        use_default_metadata: bool = True,
        call_custom_metadata: Optional[Union[MetadataFn, Sequence[MetadataFn]]] = None,
    ) -> List[Document]:
        """
        Load a SPECIFIC file (single-file) and enrich with metadata, like in load_documents().
        Returns a list of Document (e.g., PDF → 1 per page, DOCX paragraphs → many, etc.).
        """
        p = Path(file_path).resolve()
        docs: List[Document] = []

        if not p.exists() or not p.is_file():
            logger.error("[intergraxDocumentsLoader] File not found: %s", p)
            return docs

        # size / exclude / extension
        if not self._is_within_limits(p):
            return docs

        ext = p.suffix.lower()
        if ext not in self._allowed_exts:
            logger.warning("[intergraxDocumentsLoader] Unsupported extension for single-file load: %s", p)
            return docs

        # callbacks → list
        callbacks: List[MetadataFn] = []
        if call_custom_metadata:
            callbacks = (
                list(call_custom_metadata)
                if isinstance(call_custom_metadata, (list, tuple))
                else [call_custom_metadata]
            )

        try:
            loader_factory = self._extensions_map.get(ext)
            if loader_factory is None:
                logger.warning("[intergraxDocumentsLoader] No loader for extension: %s", ext)
                return docs

            loader = loader_factory(str(p))
            loaded = loader.load()  # typically List[Document]
            if not loaded:
                return docs

            parent_id = self._stable_parent_id(p)

            for d in loaded:
                if use_default_metadata:
                    d.metadata.setdefault("source_path", str(p))
                    d.metadata.setdefault("source_name", p.name)
                    d.metadata.setdefault("ext", ext)
                    if "page" in d.metadata and "page_index" not in d.metadata:
                        d.metadata["page_index"] = d.metadata["page"]
                    d.metadata.setdefault("parent_id", parent_id)

                for cb in callbacks:
                    try:
                        extra = cb(d, p)
                        if isinstance(extra, dict):
                            d.metadata.update({k: v for k, v in extra.items() if v is not None})
                    except Exception as cb_e:
                        logger.exception("[intergraxDocumentsLoader] Callback error for %s: %s", p, cb_e)

            docs.extend(loaded)
            return docs

        except Exception as e:
            logger.exception("[intergraxDocumentsLoader] Error while loading file %s: %s", p, e)
            return docs


    def load_documents(
        self,
        directory_path: str,
        *,
        use_default_metadata: bool = True,
        call_custom_metadata: Optional[Union[MetadataFn, Sequence[MetadataFn]]] = None
    ) -> List[Document]:
        """
        Scans a directory according to file_patterns/exclusions/limits and
        delegates each file to load_document(...).
        """
        if self._verbose:
            logger.info("[intergraxDocumentsLoader] Loading documents from %s", directory_path)

        docs: List[Document] = []
        root = Path(directory_path).resolve()
        if not root.exists():
            logger.error("[intergraxDocumentsLoader] Directory not found: %s", root)
            return docs

        # Gather candidates by patterns
        all_files: List[Path] = []
        for pattern in self._file_patterns:
            for f in root.glob(pattern):
                try:
                    if not self._follow_symlinks and f.is_symlink():
                        continue
                    all_files.append(f)
                except OSError:
                    continue

        # Filtering: file, extension, exclude, size
        candidate_files: List[Path] = []
        for f in all_files:
            try:
                if not f.is_file():
                    continue
                if self._is_excluded(f, root):
                    continue
                if f.suffix.lower() not in self._allowed_exts:
                    continue
                if not self._is_within_limits(f):
                    continue
                candidate_files.append(f)
            except OSError:
                continue

        # File count limit
        if self._max_files is not None and len(candidate_files) > self._max_files:
            if self._verbose:
                logger.warning(
                    "[intergraxDocumentsLoader] Too many files (%d). Truncating to %d.",
                    len(candidate_files),
                    self._max_files,
                )
            candidate_files = candidate_files[: self._max_files]

        # Progress bar + delegation to single-file loader
        with tqdm(
            desc=f"Loading files from {directory_path}",
            unit="file",
            leave=False,
            total=len(candidate_files),
            disable=not self._verbose,
        ) as pbar:
            for file in candidate_files:
                try:
                    file_docs = self.load_document(
                        str(file),
                        use_default_metadata=use_default_metadata,
                        call_custom_metadata=call_custom_metadata,
                    )
                    if file_docs:
                        docs.extend(file_docs)
                except Exception as e:
                    logger.exception("[intergraxDocumentsLoader] Error while loading file %s: %s", file, e)
                finally:
                    pbar.update(1)

        if self._verbose:
            logger.info("[intergraxDocumentsLoader] Done. Loaded documents: %d", len(docs))

        return docs



class DocxParagraphLoader:
    """
    DOCX → list of Documents (1 per paragraph or heading).
    Returns List[langchain_core.documents.Document], compatible with your load_documents().
    """
    def __init__(self, path: str, mode: str = "paragraphs"):  # "paragraphs" | "headings"
        if docx is None:
            raise ImportError("python-docx is required for DocxParagraphLoader (pip install python-docx)")
        self.path = path
        self.mode = mode

    def _is_heading(self, para) -> tuple[bool, int]:
        """Returns (is_heading, level 1..9 or 0)."""
        style = getattr(para.style, "name", "") or ""
        if not style:
            return (False, 0)
        s = style.lower()
        if s.startswith("heading"):
            # e.g., 'Heading 1'.. 'Heading 9'
            for i in range(1, 10):
                if s == f"heading {i}":
                    return (True, i)
            return (True, 1)
        return (False, 0)

    def load(self):
        from langchain_core.documents import Document
        d = docx.Document(self.path)
        items = []
        heading_stack: list[str] = []   # ["H1", "H2", ...]
        section_ix = 0
        para_ix = 0

        for p in d.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue

            is_head, level = self._is_heading(p)
            if is_head:
                # update heading path
                heading_stack = heading_stack[:max(level - 1, 0)]
                heading_stack.append(text)
                section_ix += 1

                if self.mode == "headings":
                    meta = {
                        "doc_type": "docx",
                        "source_path": self.path,
                        "source_name": str(self.path).split("/")[-1],
                        "section_ix": section_ix,
                        "heading_path": " / ".join(heading_stack),
                        "is_heading": True,
                    }
                    items.append(Document(page_content=text, metadata=meta))
                continue

            # paragraphs mode: each paragraph as a separate document
            if self.mode == "paragraphs":
                para_ix += 1
                meta = {
                    "doc_type": "docx",
                    "source_path": self.path,
                    "source_name": str(self.path).split("/")[-1],
                    "section_ix": section_ix,  # 0 if paragraphs before the first Hx
                    "para_ix": para_ix,
                    "heading_path": " / ".join(heading_stack),
                    "is_heading": False,
                }
                items.append(Document(page_content=text, metadata=meta))

        return items



class PdfSmartLoader:
    """
    PDF → list of Document(s), with OCR fallback ONLY for pages that are empty after text extraction.
    - Main text source: PyMuPDFLoader (langchain community)
    - Fallback: rasterization to bitmap and pytesseract.image_to_string(...)
    """
    def __init__(
        self,
        path: str,
        *,
        enable_ocr: bool = False,
        ocr_lang: str = "eng",
        ocr_dpi: int = 200,
        ocr_psm: int | None = None,
        ocr_oem: int | None = None,
        ocr_max_pages: int | None = None,   # hard-cap on number of pages for OCR (None = no limit)
    ):
        from langchain_community.document_loaders import PyMuPDFLoader
        self.path = path
        self.enable_ocr = bool(enable_ocr)
        self.ocr_lang = ocr_lang
        self.ocr_dpi = int(ocr_dpi)
        self.ocr_psm = ocr_psm
        self.ocr_oem = ocr_oem
        self.ocr_max_pages = ocr_max_pages
        self._base = PyMuPDFLoader(path)

    def _ocr_page(self, page) -> str:
        """Rasterizes a page and performs OCR. Returns text (may be empty)."""
        if not (fitz and pytesseract and Image):
            return ""
        pix = page.get_pixmap(dpi=self.ocr_dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        cfg_parts = []
        if self.ocr_psm is not None:
            cfg_parts.append(f"--psm {int(self.ocr_psm)}")
        if self.ocr_oem is not None:
            cfg_parts.append(f"--oem {int(self.ocr_oem)}")
        config = " ".join(cfg_parts) if cfg_parts else None
        try:
            return pytesseract.image_to_string(img, lang=self.ocr_lang, config=config) or ""
        except Exception:
            return ""

    def load(self) -> list[Document]:
        # 1) first do the “normal” parsing
        docs = self._base.load()  # usually each Document = one page (metadata['page'])
        if not docs or not self.enable_ocr:
            return docs

        # 2) identify empty pages and do OCR only there
        # open the PDF with PyMuPDF once (if available)
        pdf = None
        if fitz:
            try:
                pdf = fitz.open(self.path)
            except Exception:
                pdf = None

        ocr_done = 0
        for d in docs:
            content = (d.page_content or "").strip()
            if content:
                continue
            if self.ocr_max_pages is not None and ocr_done >= self.ocr_max_pages:
                break
            # find page index (PyMuPDFLoader adds 'page' → copied below as 'page_index' in intergrax loader)
            pidx = d.metadata.get("page") or d.metadata.get("page_index")
            if pidx is None or pdf is None:
                continue
            try:
                page = pdf.load_page(int(pidx))
            except Exception:
                continue
            ocr_text = (self._ocr_page(page) or "").strip()
            if ocr_text:
                d.page_content = ocr_text
                # mark that it came from OCR (useful for debugging/metrics)
                md = dict(d.metadata or {})
                md["ocr"] = True
                md["ocr_lang"] = self.ocr_lang
                md["ocr_dpi"] = self.ocr_dpi
                d.metadata = md
                ocr_done += 1

        if pdf is not None:
            try:
                pdf.close()
            except Exception:
                pass
        return docs
    

class ExcelSmartLoader:
    """
    Excel/CSV → list of Documents.
    Modes:
      - rows:    1 doc = 1 row (page_content = JSON; best for RAG)
      - sheets:  1 doc = 1 sheet as a Markdown table (smaller files / overview)
      - markdown:1 doc = 1 row as readable text "col: value"
    Supports: .xlsx, .xls, .csv, .tsv
    """

    def __init__(
        self,
        path: str,
        *,
        mode: str = "rows",            # "rows" | "sheets" | "markdown"
        header: int = 0,               # header row (as in pandas), None = no headers
        sheet: str | int | None = None,# sheet name/index; None => all
        na_filter: bool = True,        # drop empty rows (all NaN/empty)
        max_rows_per_sheet: int | None = None,  # limit for huge files
        encoding: str | None = None,   # for CSV/TSV
        delimiter: str | None = None,  # for CSV/TSV; None => auto (, / \t)
    ):
        if pd is None:
            raise ImportError("pandas is required for ExcelSmartLoader (pip install pandas openpyxl)")

        self.path = path
        self.mode = mode
        self.header = header
        self.sheet = sheet
        self.na_filter = na_filter
        self.max_rows_per_sheet = max_rows_per_sheet
        self.encoding = encoding
        self.delimiter = delimiter

    def _is_excel(self) -> bool:
        low = self.path.lower()
        return low.endswith(".xlsx") or low.endswith(".xls")

    def _is_tsv(self) -> bool:
        return self.path.lower().endswith(".tsv")

    def _read_excel(self) -> dict[str, DataFrame]:
        # sheet_name=None => dict of DataFrames
        kwargs = dict(sheet_name=self.sheet if self.sheet is not None else None, header=self.header, engine=None)
        # pandas will automatically choose openpyxl/xlrd if available
        dfs = pd.read_excel(self.path, **kwargs)
        if isinstance(dfs, DataFrame):
            # single sheet → wrap in dict
            return {"Sheet1": dfs}
        return dfs  # {sheet_name: df}

    def _read_csv_like(self) -> dict[str, DataFrame]:
        # Treat CSV/TSV as "one sheet"
        sep = self.delimiter
        if sep is None:
            sep = "\t" if self._is_tsv() else ","  # simple heuristic choice
        df = pd.read_csv(self.path, sep=sep, header=self.header, encoding=self.encoding)
        return {"csv": df}

    @staticmethod
    def _dtype_map(df: DataFrame) -> dict:
        return {str(c): str(t) for c, t in df.dtypes.items()}

    @staticmethod
    def _row_to_json(df: DataFrame, idx: int) -> str:
        rec = df.iloc[idx].to_dict()
        # safe JSON (strings, numbers, dates → str)
        def _safe(v):
            if pd.isna(v):
                return None
            if hasattr(v, "isoformat"):
                try:
                    return v.isoformat()
                except Exception:
                    return str(v)
            return v if isinstance(v, (int, float, bool, str)) else str(v)
        rec = {str(k): _safe(v) for k, v in rec.items()}
        return json.dumps(rec, ensure_ascii=False)

    @staticmethod
    def _row_to_markdown(df: DataFrame, idx: int) -> str:
        rec = df.iloc[idx].to_dict()
        parts = []
        for k, v in rec.items():
            if pd.isna(v):
                continue
            parts.append(f"- **{k}**: {v}")
        return "\n".join(parts) if parts else "- (empty row)"

    @staticmethod
    def _df_to_markdown(df: DataFrame, max_rows: int | None = None) -> str:
        _df = df if max_rows is None else df.head(max_rows)
        try:
            return _df.to_markdown(index=False)
        except Exception:
            # fallback (without tabulate dependency)
            header = " | ".join(map(str, _df.columns))
            sep = " | ".join(["---"] * len(_df.columns))
            rows = [" | ".join(map(lambda x: "" if pd.isna(x) else str(x), r)) for _, r in _df.iterrows()]
            return "\n".join([header, sep] + rows)

    def load(self) -> list[Document]:
        docs: list[Document] = []
        is_excel = self._is_excel()

        sheets = self._read_excel() if is_excel else self._read_csv_like()

        for sname, df in sheets.items():
            if self.na_filter:
                df = df.dropna(how="all")
            n_rows, n_cols = int(df.shape[0]), int(df.shape[1])
            headers = [str(c) for c in df.columns]
            dtype_map = self._dtype_map(df)

            # hard cap (huge files)
            row_cap = self.max_rows_per_sheet if self.max_rows_per_sheet is not None else n_rows

            if self.mode == "sheets":
                content = self._df_to_markdown(df, max_rows=row_cap)
                docs.append(Document(
                    page_content=content,
                    metadata={
                        "source_name": self.path.split("/")[-1],
                        "source_path": self.path,
                        "ext": ".xlsx" if is_excel else ".csv",
                        "sheet_name": sname,
                        "n_rows": n_rows,
                        "n_cols": n_cols,
                        # list/dict → JSON string:
                        "headers_json": json.dumps(headers, ensure_ascii=False),
                        "dtype_map_json": json.dumps(dtype_map, ensure_ascii=False),
                        "excel_mode": "sheets",
                    }
                ))
                continue

            # rows / markdown
            max_i = min(row_cap, n_rows)
            for i in range(max_i):
                content = (
                    self._row_to_json(df, i) if self.mode == "rows"
                    else self._row_to_markdown(df, i)
                )
                docs.append(Document(
                    page_content=content,
                    metadata={
                        "source_name": self.path.split("/")[-1],
                        "source_path": self.path,
                        "ext": ".xlsx" if is_excel else ".csv",
                        "sheet_name": sname,
                        "row_ix": i,
                        "n_rows": n_rows,
                        "n_cols": n_cols,
                        # list/dict → JSON string:
                        "headers_json": json.dumps(headers, ensure_ascii=False),
                        "dtype_map_json": json.dumps(dtype_map, ensure_ascii=False),
                        "excel_mode": self.mode,
                    }
                ))

        return docs
    

class ImageSmartLoader:
    """
    Universal image loader: JPG, PNG, TIFF, BMP, WEBP, HEIC/HEIF.
    Modes:
      - OCR:     extract visible text with Tesseract
      - Caption: call your framework LLM adapter (e.g., Ollama) to describe the image
      - Both:    combine caption + OCR (with a joiner)
    Always returns 1 Document per image with clear provenance metadata.
    """

    def __init__(
        self,
        path: str,
        *,
        ocr_lang: str = "eng",
        ocr_psm: int | None = None,
        ocr_oem: int | None = None,
        extract_exif: bool = True,
        max_image_dim: int | None = None,  # e.g., 2000 – downscale if larger
        # NEW:
        text_mode: Literal["ocr", "caption", "both"] = "both",
        caption_llm: Optional[LLMAdapter] = None,
        both_joiner: str = "\n\n---\n\n",
    ):
        self.path = path
        self.ocr_lang = ocr_lang
        self.ocr_psm = ocr_psm
        self.ocr_oem = ocr_oem
        self.extract_exif = bool(extract_exif)
        self.max_image_dim = max_image_dim

        self.text_mode = text_mode
        self.caption_llm = caption_llm
        self.both_joiner = both_joiner

    # ---------- helpers ----------
    def _resize_if_needed(self, img: Image) -> Image:
        if self.max_image_dim is None:
            return img
        w, h = img.size
        if max(w, h) <= self.max_image_dim:
            return img
        ratio = self.max_image_dim / float(max(w, h))
        new_size = (int(w * ratio), int(h * ratio))
        return img.resize(new_size)

    def _ocr(self, img: Image) -> str:
        if pytesseract is None:
            return ""
        cfg_parts = []
        if self.ocr_psm is not None:
            cfg_parts.append(f"--psm {int(self.ocr_psm)}")
        if self.ocr_oem is not None:
            cfg_parts.append(f"--oem {int(self.ocr_oem)}")
        config = " ".join(cfg_parts) if cfg_parts else None
        try:
            return pytesseract.image_to_string(img, lang=self.ocr_lang, config=config) or ""
        except Exception:
            return ""

    def _infer_ollama_model(self) -> Optional[str]:
        """
        Try to infer model name from the LangChainOllamaAdapter.
        - Prefer adapter.defaults.get("model")
        - Fallbacks are possible (chat.model), else None
        """
        if not isinstance(self.caption_llm, LangChainOllamaAdapter):
            return None
        defaults = getattr(self.caption_llm, "defaults", {}) or {}
        model = defaults.get("model")
        if model:
            return model
        chat = getattr(self.caption_llm, "chat", None)
        if chat is not None:
            for attr in ("model", "model_name", "model_id"):
                if hasattr(chat, attr):
                    try:
                        val = getattr(chat, attr)
                        if isinstance(val, str) and val.strip():
                            return val.strip()
                    except Exception:
                        pass
            for attr in ("kwargs", "config", "client"):
                try:
                    obj = getattr(chat, attr, None)
                    if isinstance(obj, dict):
                        for k in ("model", "model_name", "model_id"):
                            if isinstance(obj.get(k), str) and obj[k].strip():
                                return obj[k].strip()
                except Exception:
                    pass
        return None

    def _caption_via_ollama(self, img_path: str) -> str:
        """
        Vision caption bridge for Ollama (local REST).
        Uses model inferred from the adapter. Endpoint defaults to localhost.
        """
        from intergrax.multimedia.images_loader import transcribe_image

        model = self._infer_ollama_model() or "llava-llama3:latest"
        prompt = "Describe the image in detail."


        resp = transcribe_image(
            prompt=prompt, 
            model=model,
            image_path=img_path,
        )

        return resp

    def _caption_via_adapter(self, img_path: str) -> str:
        """
        Generic bridge:
        - If adapter exposes describe_image(path) → use it.
        - Else if it's LangChainOllamaAdapter → use REST vision bridge.
        - Else raise (you can extend here for OpenAI/Gemini Vision).
        """
        if self.caption_llm is None:
            return ""
        # 1) Native helper, if adapter ją posiada
        if hasattr(self.caption_llm, "describe_image"):
            try:
                txt = self.caption_llm.describe_image(img_path)
                return (txt or "").strip()
            except Exception as e:
                raise RuntimeError(f"LLMAdapter.describe_image failed: {e}")
        # 2) Ollama vision fallback
        if isinstance(self.caption_llm, LangChainOllamaAdapter):
            return self._caption_via_ollama(img_path)
        # 3) Not supported yet
        raise ValueError("Captioning supported for adapters exposing describe_image(...) or LangChainOllamaAdapter (vision).")

    def _exif_dict(self, img: Image) -> dict:
        out = {}
        if not (self.extract_exif and ExifTags and hasattr(img, "_getexif")):
            return out
        try:
            exif_raw = img._getexif() or {}
            for tag, value in exif_raw.items():
                tag_name = ExifTags.TAGS.get(tag, str(tag))
                out[tag_name] = str(value)
        except Exception:
            pass
        return out

    # ---------- main ----------
    def load(self) -> list[Document]:
        if Image is None:
            raise ImportError("Pillow (PIL) is required for ImageSmartLoader")

        img = Image.open(self.path)
        img = self._resize_if_needed(img)
        width, height = img.size
        dpi = img.info.get("dpi", None)
        exif = self._exif_dict(img)

        # Decide which mechanisms to run
        run_ocr = (self.text_mode in ("ocr", "both")) and (pytesseract is not None)
        run_caption = (self.text_mode in ("caption", "both")) and (self.caption_llm is not None)

        ocr_text = self._ocr(img) if run_ocr else ""
        caption_text = self._caption_via_adapter(self.path) if run_caption else ""

        # Assemble content
        if self.text_mode == "ocr":
            content = (ocr_text or "").strip() or "(No visible text detected.)"
        elif self.text_mode == "caption":
            content = (caption_text or "").strip() or "(No caption generated.)"
        else:
            left = (caption_text or "").strip()
            right = (ocr_text or "").strip()
            if left and right:
                content = f"{left}{self.both_joiner}{right}"
            elif left:
                content = left
            elif right:
                content = right
            else:
                content = "(No caption nor OCR text produced.)"

        # Metadata
        meta = {
            "source_name": os.path.basename(self.path),
            "source_path": self.path,
            "format": img.format,
            "width": width,
            "height": height,
            "dpi": (dpi[0] if isinstance(dpi, tuple) and len(dpi) > 0 else dpi),
            "exif_json": json.dumps(exif, ensure_ascii=False) if exif else None,

            # Provenance & modes
            "image_text_mode": self.text_mode,                 # "ocr" | "caption" | "both"
            "ocr_lang": self.ocr_lang if (self.text_mode in ("ocr", "both")) else None,
            "caption_llm": type(self.caption_llm).__name__ if (self.text_mode in ("caption", "both") and self.caption_llm) else None,
            "caption_model_inferred": (self._infer_ollama_model() or "llava-llama3:latest")
                if (self.text_mode in ("caption", "both") and isinstance(self.caption_llm, LangChainOllamaAdapter))
                else None,
        }

        return [Document(page_content=content, metadata=meta)]


class VideoSmartLoader:
    """
    Loads video files and converts them into a list of LangChain Documents.

    For each subtitle (VTT) segment:
      - page_content: transcript text
      - metadata: extracted frame path, mid_time_ms, video_segment_id, etc.

    If a .vtt transcript file is missing, it can automatically generate one
    using Whisper via transcribe_to_vtt().

    It also extracts key video frames using extract_and_frames_and_metadata().
    """

    def __init__(
        self,
        path: str,
        *,
        out_dir: str | None = None,
        frames_subdir: str = "frames",
        meta_subdir: str = "video_meta",
        transcribe_if_missing: bool = True,
        whisper_model_size: str = "base",
        whisper_language: str | None = None,
        frame_target_height: int = 350,
    ):
        """
        Args:
            path: Path to the video file.
            out_dir: Optional directory to save extracted data.
            frames_subdir: Subfolder for extracted frames.
            meta_subdir: Subfolder for metadata.
            transcribe_if_missing: Whether to generate .vtt if missing.
            whisper_model_size: Whisper model name ('tiny', 'base', etc.).
            whisper_language: Optional language code (e.g. 'en', 'pl').
            frame_target_height: Frame resize height while keeping aspect ratio.
        """
        self.path = str(path)
        self._p = Path(path).resolve()

        # Determine output directories
        self.out_root = Path(out_dir) if out_dir else self._p.parent
        self.frames_dir = self.out_root / frames_subdir
        self.meta_dir = self.out_root / meta_subdir

        self.transcribe_if_missing = bool(transcribe_if_missing)
        self.whisper_model_size = whisper_model_size
        self.whisper_language = whisper_language
        self.frame_target_height = int(frame_target_height)

        self.frames_dir.mkdir(parents=True, exist_ok=True)
        self.meta_dir.mkdir(parents=True, exist_ok=True)

    def _ensure_vtt(self) -> str:
        """
        Checks if a .vtt transcript exists next to the video file.
        If not, optionally generates it using Whisper.
        Returns the absolute path to the .vtt file.
        """
        vtt_path = self._p.with_suffix(".vtt")
        if vtt_path.exists():
            return str(vtt_path)

        if not self.transcribe_if_missing:
            raise FileNotFoundError(f"Missing transcript: {vtt_path}")

        # Generate transcript using your component (Whisper)
        vtt_path = transcribe_to_vtt(
            input_media_path=str(self._p),
            output_vtt_path=str(vtt_path),
            model_size=self.whisper_model_size,
            language=self.whisper_language,
        )
        return str(vtt_path)

    def load(self) -> list[Document]:
        """
        Extracts transcript and representative frames per subtitle segment.
        Returns a list of LangChain Document objects.
        """
        # 1. Ensure transcript exists
        vtt_path = self._ensure_vtt()

        # 2. Extract frames and metadata using your helper
        metas = extract_frames_and_metadata(
            path_to_video=str(self._p),
            path_to_transcript=str(vtt_path),
            path_to_save_extracted_frames=str(self.frames_dir),
            path_to_save_metadatas=str(self.meta_dir),
        )

        docs: list[Document] = []
        for m in metas:
            transcript = (m.get("transcript") or "").strip()
            if not transcript:
                continue

            metadata = {
                "doc_type": "video",
                "video_path": m.get("video_path") or str(self._p),
                "video_segment_id": m.get("video_segment_id"),
                "mid_time_ms": m.get("mid_time_ms"),
                "extracted_frame_path": m.get("extracted_frame_path"),
                "transcript_source": "vtt",
            }

            # Include optional timing/frame info if available
            for k in ("start_ms", "end_ms", "start", "end", "frame_index","duration_ms"):
                if k in m and m[k] is not None:
                    metadata[k] = m[k]

            docs.append(Document(page_content=transcript, metadata=metadata))

        return docs



class AudioSmartLoader:
    """
    Smart loader for audio files and YouTube audio sources.

    Capabilities:
    - Accepts either a local audio file or a YouTube URL.
    - Downloads the audio (if YouTube URL) using yt_dlp.
    - Transcribes or translates audio using Whisper (translate_audio).
    - Returns a list of LangChain Document objects with metadata per segment.
    """

    def __init__(
        self,
        path: str,
        *,
        out_dir: str | Path | None = None,
        audio_format: str = "mp3",
        whisper_model: str = "medium",
        whisper_language: str = "en",
        translate: bool = True,
    ):
        """
        Args:
            path: Local audio path or YouTube URL.
            out_dir: Directory for downloaded audio files (if YouTube).
            audio_format: Desired output format (mp3, wav, flac, etc.).
            whisper_model: Whisper model name (tiny, base, small, medium, large).
            whisper_language: Language code for transcription/translation.
            translate: If True, translates speech to English (Whisper task="translate").
        """
        self.path = path
        self.out_dir = Path(out_dir or "./audio_downloads")
        self.audio_format = audio_format
        self.whisper_model = whisper_model
        self.whisper_language = whisper_language
        self.translate = translate

    def load(self) -> List[Document]:
        """
        Executes the end-to-end audio pipeline:
        - If path is a YouTube URL → download audio.
        - Run Whisper transcription/translation.
        - Return list of LangChain Documents with metadata.
        """
        audio_path = self._ensure_audio_file()
        transcription = self._transcribe_audio(audio_path)

        # Whisper output example:
        # {
        #   "text": "full transcription",
        #   "segments": [
        #       {"id":0,"start":0.0,"end":3.1,"text":"Hello world"},
        #       ...
        #   ]
        # }

        results = transcription.get("segments", [])
        if not results:
            # fallback: single document with whole text
            return [
                Document(
                    page_content=transcription.get("text", "").strip(),
                    metadata={
                        "source_path": str(audio_path),
                        "source_type": "audio",
                        "language": self.whisper_language,
                        "whisper_model": self.whisper_model,
                    },
                )
            ]

        # Build documents for each segment
        docs: List[Document] = []
        for seg in results:
            seg_text = seg.get("text", "").strip()
            seg_start = float(seg.get("start", 0))
            seg_end = float(seg.get("end", 0))
            duration = seg_end - seg_start

            metadata = {
                "source_path": str(audio_path),
                "source_type": "audio",
                "segment_id": seg.get("id"),
                "start_s": seg_start,
                "end_s": seg_end,
                "duration_s": duration,
                "whisper_model": self.whisper_model,
                "language": self.whisper_language,
                "translated": self.translate,
            }

            docs.append(Document(page_content=seg_text, metadata=metadata))

        return docs

    # ---------------------------------------------------------
    # Internal helpers
    # ---------------------------------------------------------
    def _ensure_audio_file(self) -> Path:
        """If given a YouTube URL, downloads the audio file."""
        if self.path.startswith("http://") or self.path.startswith("https://"):
            # YouTube link → download audio
            print(f"[AudioSmartLoader] Downloading audio from YouTube: {self.path}")
            audio_path = yt_download_audio(
                youtube_url=self.path,
                out_dir=self.out_dir,
                audio_format=self.audio_format,
            )
            print(f"[AudioSmartLoader] Audio downloaded: {audio_path}")
            return audio_path

        # Local file
        audio_path = Path(self.path)
        if not audio_path.exists():
            raise FileNotFoundError(f"Audio file not found: {audio_path}")
        return audio_path

    def _transcribe_audio(self, audio_path: Path) -> dict:
        """Transcribes or translates the given audio file using Whisper."""
        print(f"[AudioSmartLoader] Transcribing audio: {audio_path}")
        try:
            result = translate_audio(
                str(audio_path),
                model=self.whisper_model,
                language=self.whisper_language,
            )
            return result
        except Exception as e:
            raise RuntimeError(f"Failed to transcribe audio {audio_path}: {e}")
