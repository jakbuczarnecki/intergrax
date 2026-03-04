# © Artur Czarnecki. All rights reserved.
# Intergrax framework – proprietary and confidential.
# Use, modification, or distribution without written permission is prohibited.

from __future__ import annotations

from typing import Literal, Sequence
import docx
from langchain_core.documents import Document
from langchain_community.document_loaders import Docx2txtLoader

from intergrax.rag.document_loaders.config.document_loader_config import (
    DEFAULT_BUILTIN_HANDLER_CONFIDENCE,
)
from intergrax.rag.document_loaders.contracts.base_document_handler import (
    BaseDocumentHandler,
)

EXTRACTION_STRATEGY = Literal["auto", "fulltext", "paragraphs", "headings"]

class DocSmartDocumentHandler(BaseDocumentHandler):

    def __init__(self, extraction_strategy: EXTRACTION_STRATEGY = "auto") -> None:
        self._extraction_strategy = extraction_strategy

    def supports(self, source: str) -> bool:
        s = source.lower()
        return s.endswith(".docx") or s.endswith(".doc")

    def confidence(self, source: str) -> float:
        return DEFAULT_BUILTIN_HANDLER_CONFIDENCE

    def load(self, source: str) -> Sequence[Document]:

        strategy = self._extraction_strategy

        if strategy == "auto":
            strategy = "fulltext"

        if strategy == "fulltext":
            loader = Docx2txtLoader(source)
            return loader.load()

        if strategy in ("paragraphs", "headings"):
            loader = DocxParagraphLoader(source, mode=strategy)
            return loader.load()

        raise RuntimeError("Invalid extraction strategy state")


class DocxParagraphLoader:
        """
        DOCX → list of Documents (1 per paragraph or heading).
        Returns List[langchain_core.documents.Document], compatible with your load_documents().
        """
        def __init__(self, path: str, mode: EXTRACTION_STRATEGY = "paragraphs"):
            if docx is None:
                raise ImportError("python-docx is required for DocxParagraphLoader (pip install python-docx)")
            self.path = path
            self.mode = mode

        def _is_heading(self, para) -> tuple[bool, int]:
            """Returns (is_heading, level 1..9 or 0)."""
            style = getattr(para.style, "name", "") or ""
            if not style:
                return (False, 0)
            s = style.lower()
            if s.startswith("heading"):
                # e.g., 'Heading 1'.. 'Heading 9'
                for i in range(1, 10):
                    if s == f"heading {i}":
                        return (True, i)
                return (True, 1)
            return (False, 0)

        def load(self):
            from langchain_core.documents import Document
            d = docx.Document(self.path)
            items = []
            heading_stack: list[str] = []   # ["H1", "H2", ...]
            section_ix = 0
            para_ix = 0

            for p in d.paragraphs:
                text = (p.text or "").strip()
                if not text:
                    continue

                is_head, level = self._is_heading(p)
                if is_head:
                    # update heading path
                    heading_stack = heading_stack[:max(level - 1, 0)]
                    heading_stack.append(text)
                    section_ix += 1

                    if self.mode == "headings":
                        meta = {
                            "doc_type": "docx",
                            "source_path": self.path,
                            "source_name": str(self.path).split("/")[-1],
                            "section_ix": section_ix,
                            "heading_path": " / ".join(heading_stack),
                            "is_heading": True,
                        }
                        items.append(Document(page_content=text, metadata=meta))
                    continue

                # paragraphs mode: each paragraph as a separate document
                if self.mode == "paragraphs":
                    para_ix += 1
                    meta = {
                        "doc_type": "docx",
                        "source_path": self.path,
                        "source_name": str(self.path).split("/")[-1],
                        "section_ix": section_ix,  # 0 if paragraphs before the first Hx
                        "para_ix": para_ix,
                        "heading_path": " / ".join(heading_stack),
                        "is_heading": False,
                    }
                    items.append(Document(page_content=text, metadata=meta))

            return items