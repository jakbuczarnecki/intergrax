# © Artur Czarnecki. All rights reserved.
# Intergrax framework – proprietary and confidential.
# Use, modification, or distribution without written permission is prohibited.

from __future__ import annotations

from typing import Literal, Sequence

import docx
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.documents import Document

from intergrax.rag.document_loaders.contracts.base_document_parser import BaseDocumentParser
from intergrax.rag.document_loaders.contracts.metadata_contract import build_loader_metadata

EXTRACTION_STRATEGY = Literal["auto", "fulltext", "paragraphs", "headings"]

class DocSmartParser(BaseDocumentParser):

    def __init__(self, strategy: EXTRACTION_STRATEGY):
        self._strategy = strategy

    @classmethod
    def parser_id(cls) -> str:
        return "doc_smart"

    def is_available(self) -> bool:
        return True

    def load(self, source: str) -> Sequence[Document]:

        strategy = self._strategy

        if strategy == "auto":
            strategy = "fulltext"

        if strategy == "fulltext":
            loader = Docx2txtLoader(source)
            docs = loader.load()

        elif strategy in ("paragraphs", "headings"):
            loader = DocxParagraphLoader(source, mode=strategy)
            docs = loader.load()

        else:
            raise RuntimeError("Invalid extraction strategy state")

        result: list[Document] = []

        for i, d in enumerate(docs):

            metadata = build_loader_metadata(
                source=source,
                parser=self.parser_id(),
                position=i,
            )

            metadata.update(d.metadata or {})

            result.append(
                Document(
                    page_content=d.page_content,
                    metadata=metadata,
                )
            )

        return result


class DocxParagraphLoader:

    def __init__(self, path: str, mode: EXTRACTION_STRATEGY = "paragraphs"):
        if docx is None:
            raise ImportError(
                "python-docx is required for DocxParagraphLoader (pip install python-docx)"
            )
        self.path = path
        self.mode = mode

    def _is_heading(self, para) -> tuple[bool, int]:

        style = getattr(para.style, "name", "") or ""
        if not style:
            return (False, 0)

        s = style.lower()

        if s.startswith("heading"):
            for i in range(1, 10):
                if s == f"heading {i}":
                    return (True, i)
            return (True, 1)

        return (False, 0)

    def load(self):

        d = docx.Document(self.path)

        items = []
        heading_stack: list[str] = []
        section_ix = 0
        para_ix = 0

        for p in d.paragraphs:

            text = (p.text or "").strip()

            if not text:
                continue

            is_head, level = self._is_heading(p)

            if is_head:

                heading_stack = heading_stack[:max(level - 1, 0)]
                heading_stack.append(text)

                section_ix += 1

                if self.mode == "headings":

                    meta = {
                        "doc_type": "docx",
                        "source_path": self.path,
                        "source_name": str(self.path).split("/")[-1],
                        "section_ix": section_ix,
                        "heading_path": " / ".join(heading_stack),
                        "is_heading": True,
                    }

                    items.append(Document(page_content=text, metadata=meta))

                continue

            if self.mode == "paragraphs":

                para_ix += 1

                meta = {
                    "doc_type": "docx",
                    "source_path": self.path,
                    "source_name": str(self.path).split("/")[-1],
                    "section_ix": section_ix,
                    "para_ix": para_ix,
                    "heading_path": " / ".join(heading_stack),
                    "is_heading": False,
                }

                items.append(Document(page_content=text, metadata=meta))

        return items